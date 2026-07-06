"""web.docx_fill の単体テスト（LLM 非依存）。

園の実 Word 様式（`web/templates/*.docx`）へ確定 entry を流し込む presentation 層の描画を検証する。
型検査は harness の責務なので、ここは「値が正しいセルに入るか／未対応は落ちるか」だけを見る（§5）。
chohyo_pdf のテスト（tests/test_chohyo_pdf.py）と対になる。
"""

from __future__ import annotations

import io

import docx
import pytest
from docx.oxml.ns import qn

from hoiku_agent.schemas import (
    AgeBand,
    ChildRecord,
    ClassMonthlyPlan,
    DevelopmentNote,
    FiveDomains,
    IndividualGoal,
    MonthlyEducationNote,
    MonthlyPlan,
    NurseryRecord,
    ThreeViewpoint,
)
from hoiku_agent.web.docx_fill import (
    _resize_table_data_rows,
    fill_docx,
    supported_kinds,
)


def _nursery() -> dict:
    return NurseryRecord(
        fiscal_year="2026",
        age_band=AgeBand.三から五歳,
        child_id="かなたくん",
        final_year_focus="共通の目的に向かって活動を楽しむ",
        individual_focus="友だちと協力してやり遂げる",
        development_notes=[
            DevelopmentNote(description="鉄棒に挑戦する姿が育った", tags=[FiveDomains.健康]),
            DevelopmentNote(description="考えを伝え合うようになった", tags=[FiveDomains.言葉]),
        ],
        special_notes="就学先と見通しの持ち方を引き継ぐ",
        growth_until_final="入園当初の不安から生き生きと表現する姿へ育った",
        school_name="市立ひがし小学校",
    ).model_dump(mode="json")


def _monthly(age_band: AgeBand = AgeBand.零から二歳) -> dict:
    return MonthlyPlan(
        month="2026-07",
        age_band=age_band,
        child_id="はるとくん",
        age_months="1歳6か月",
        prev_child_state="前月は探索活動が活発になった",
        nurturing_life="夏の体調管理を丁寧に行う",
        nurturing_emotion="甘えを受けとめ安心を支える",
        education=[
            MonthlyEducationNote(
                aim="水や砂の感触を楽しむ", tags=[ThreeViewpoint.身近なものと関わり感性が育つ]
            )
        ],
        monthly_goals="感触遊びで探索意欲を満たす",
        environment_support="水遊びの動線を整える",
        evaluation_reflection="感触遊びの幅を広げられた",
    ).model_dump(mode="json")


def _class_monthly(n_goals: int, age_band: AgeBand = AgeBand.零から二歳) -> dict:
    """0–2 クラス月案（園フォーム）＝個人目標を n_goals 名ぶん持つ確定 entry。"""
    return ClassMonthlyPlan(
        month="2026-07",
        age_band=age_band,
        class_name="ひよこ組",
        monthly_goal="梅雨期も健康に過ごす",
        prev_month_state="感触遊びに集中していた",
        individual_goals=[
            IndividualGoal(
                child_id=f"こども{i}くん",
                age_months=f"1歳{i}か月",
                child_state=f"子どもの姿{i}",
                aim_support=f"ねらい・配慮{i}",
            )
            for i in range(1, n_goals + 1)
        ],
    ).model_dump(mode="json")


def _record() -> dict:
    return ChildRecord(
        period="2026-04〜2026-06",
        age_band=AgeBand.三から五歳,
        child_id="ゆいちゃん",
        development_notes=[
            DevelopmentNote(
                description="好きな遊びに自分から関わる姿が増えた", tags=[FiveDomains.健康]
            ),
            DevelopmentNote(
                description="順番を待つ場面が見られるようになった", tags=[FiveDomains.人間関係]
            ),
        ],
        overall_note="安心を土台に遊びの世界を広げている",
    ).model_dump(mode="json")


def _tables(data: bytes) -> list:
    return docx.Document(io.BytesIO(data)).tables


def test_supported_kinds_has_child_record():
    assert "child_record" in supported_kinds()


def test_fill_child_record_returns_docx_bytes():
    data = fill_docx("child_record", _record())
    assert data[:2] == b"PK"  # docx = zip
    assert len(data) > 0


def test_fill_child_record_places_content_by_domain():
    """development_notes が 5領域タグごとに「領域×子どもの姿」表の該当行へ入る。"""
    tables = _tables(fill_docx("child_record", _record()))
    matrix = next(t for t in tables if "領域" in t.rows[0].cells[0].text)
    rows = {r.cells[0].text.strip(): r.cells[-1].text for r in matrix.rows[1:]}
    assert "好きな遊びに自分から関わる姿が増えた" in rows["健康"]
    assert "順番を待つ場面が見られるようになった" in rows["人間関係"]
    # タグの無い領域は空欄のまま（様式に無いものを勝手に作らない）。
    assert rows["環境"].strip() == ""


def test_fill_child_record_fills_header_and_period():
    tables = _tables(fill_docx("child_record", _record()))
    header = next(t for t in tables if "児童名" in "".join(c.text for c in t.rows[1].cells))
    assert "ゆいちゃん" in "".join(c.text for c in header.rows[1].cells)
    assert "2026年度" in "".join(c.text for c in header.rows[0].cells)  # 対象期間から年度推定
    period_tbl = next(t for t in tables if t.rows[0].cells[0].text.strip() == "対象期間")
    assert "2026-04〜2026-06" in period_tbl.rows[0].cells[-1].text


def test_supported_kinds_has_monthly():
    assert "monthly" in supported_kinds()


def test_fill_monthly_0_2_maps_individual_to_goals_table():
    """0-2 月案フォームの「個人目標」小表に個別月案が写像される（クラス欄は温存）。"""
    tables = _tables(fill_docx("monthly", _monthly(AgeBand.零から二歳)))
    goals = next(t for t in tables if "個人目標" in t.rows[0].cells[0].text)
    row = goals.rows[2]  # 記入1行目
    assert "はるとくん" in row.cells[0].text
    assert "探索活動が活発" in row.cells[1].text  # 子どもの姿←前月の姿
    # ねらい・配慮＝今月のねらい＋養護＋教育＋環境の束（AI 内容を落とさない）。
    assert "感触遊びで探索意欲を満たす" in row.cells[2].text
    assert "養護（生命の保持）" in row.cells[2].text
    assert "水や砂の感触を楽しむ" in row.cells[2].text
    assert "感触遊びの幅を広げられた" in row.cells[3].text  # 評価・反省


def test_fill_monthly_header_has_year_and_age():
    tables = _tables(fill_docx("monthly", _monthly(AgeBand.零から二歳)))
    header = next(t for t in tables if "年度・月" in "".join(c.text for c in t.rows[0].cells))
    assert "2026年度 7月" in "".join(c.text for c in header.rows[0].cells)
    assert "0〜2歳児" in "".join(c.text for c in header.rows[1].cells)


def test_fill_monthly_3_5_has_no_goals_table_but_renders():
    """3-5 フォームは個人目標小表が無い純クラス様式＝落ちずヘッダのみ流し込む（クラス月案対応は後続）。"""
    tables = _tables(fill_docx("monthly", _monthly(AgeBand.三から五歳)))
    assert not any("個人目標" in t.rows[0].cells[0].text for t in tables)
    header = next(t for t in tables if "年度・月" in "".join(c.text for c in t.rows[0].cells))
    assert "3〜5歳児" in "".join(c.text for c in header.rows[1].cells)


def _goals_table(tables: list):
    """テンプレの「個人目標」小表を返す（r0=見出し/r1=列見出し/r2 以降=記入行）。"""
    return next(t for t in tables if "個人目標" in t.rows[0].cells[0].text)


def _row_has_full_borders(row) -> bool:
    """行の全セルが per-cell の罫線（tcBorders）を持つ＝様式どおりの枠か（add_row の崩れ検出）。"""
    return all(
        c._tc.tcPr is not None and c._tc.tcPr.find(qn("w:tcBorders")) is not None for c in row.cells
    )


def _make_table(n_rows: int, n_cols: int = 4):
    """title/header/data… を模した合成表（data_start=2 前提の記入行検証用）。"""
    doc = docx.Document()
    return doc.add_table(rows=n_rows, cols=n_cols)


def test_resize_grows_shrinks_and_matches_exact_count():
    """記入行（data_start=2 以降）をちょうど count 行にそろえる（増減・一致）。"""
    grow = _make_table(4)  # title+header+2 data
    _resize_table_data_rows(grow, 2, 5)
    assert len(grow.rows) == 2 + 5

    shrink = _make_table(4)
    _resize_table_data_rows(shrink, 2, 1)
    assert len(shrink.rows) == 2 + 1

    same = _make_table(4)
    _resize_table_data_rows(same, 2, 2)  # count==current＝no-op
    assert len(same.rows) == 2 + 2


def test_resize_no_ops_on_header_only_table():
    """記入行が1つも無い異形（title+header のみ）は複製元が無いので触らない（IndexError も出さない）。"""
    header_only = _make_table(2)  # title+header, 0 data rows
    _resize_table_data_rows(header_only, 2, 3)  # current<=0 → 早期 return
    assert len(header_only.rows) == 2


def test_supported_kinds_has_class_monthly():
    assert "class_monthly" in supported_kinds()


def test_fill_class_monthly_generates_one_row_per_child():
    """個人目標がテンプレ既定（2行）を超えても人数分の記入行を生成し、様式（罫線・4列）を保つ。"""
    tables = _tables(fill_docx("class_monthly", _class_monthly(4)))
    goals = _goals_table(tables)
    data_rows = goals.rows[2:]  # r2 以降＝記入行
    assert len(data_rows) == 4  # 4名ぶんちょうど（テンプレ2行→複製で4行）
    for i, row in enumerate(data_rows, start=1):
        assert f"こども{i}くん" in row.cells[0].text
        assert f"1歳{i}か月" in row.cells[0].text
        assert f"子どもの姿{i}" in row.cells[1].text
        assert f"ねらい・配慮{i}" in row.cells[2].text
        assert row.cells[3].text.strip() == ""  # 評価・反省は月末記入＝空欄温存
        assert len(row.cells) == 4
        assert _row_has_full_borders(row)  # 複製行も罫線を引き継ぐ（add_row の崩れを防ぐ）


def test_fill_class_monthly_renders_filled_evaluation():
    """評価・反省は通常 AI 非生成で空だが、編集フォームで記入済みなら docx にも反映する（帳票PDF と同じ）。"""
    entry = _class_monthly(1)
    entry["individual_goals"][0]["evaluation"] = "水遊びを十分に楽しめた。次月は素材を増やす"
    tables = _tables(fill_docx("class_monthly", entry))
    goals = _goals_table(tables)
    assert "水遊びを十分に楽しめた" in goals.rows[2].cells[3].text  # col3=評価・反省


def test_fill_class_monthly_removes_extra_template_rows():
    """個人目標がテンプレ既定（2行）より少なければ余った空枠を削る（人数分ちょうど）。"""
    tables = _tables(fill_docx("class_monthly", _class_monthly(1)))
    goals = _goals_table(tables)
    assert len(goals.rows) == 3  # 見出し＋列見出し＋記入1行のみ（空枠を残さない）
    assert "こども1くん" in goals.rows[2].cells[0].text


def test_fill_class_monthly_3_5_has_no_goals_table_but_renders():
    """3-5 クラス月案フォームは個人目標小表が無い＝落ちずヘッダ・グリッドのみ流し込む。"""
    tables = _tables(fill_docx("class_monthly", _class_monthly(0, AgeBand.三から五歳)))
    assert not any("個人目標" in t.rows[0].cells[0].text for t in tables)


def test_fill_monthly_shrinks_goals_table_to_single_row():
    """個別月案（1名）は記入行を1行にそろえ、テンプレの余分な空枠を残さない。"""
    tables = _tables(fill_docx("monthly", _monthly(AgeBand.零から二歳)))
    goals = _goals_table(tables)
    assert len(goals.rows) == 3  # 見出し＋列見出し＋記入1行
    assert "はるとくん" in goals.rows[2].cells[0].text


def test_supported_kinds_has_nursery_record():
    assert "nursery_record" in supported_kinds()


def test_fill_nursery_record_appends_under_labels():
    """公式様式の括弧ラベル（最終年度の重点 等）直下に内容が追記され、ラベルは残る。"""
    tables = _tables(fill_docx("nursery_record", _nursery()))
    rec = next(t for t in tables if "保育の過程" in "".join(c.text for c in t.rows[0].cells))
    joined = "\n".join(c.text for row in rec.rows for c in row.cells)
    # ガイドラベルは残す＋内容が入る（追記方式）。
    assert "（最終年度の重点）" in joined and "共通の目的に向かって活動を楽しむ" in joined
    assert "（個人の重点）" in joined and "友だちと協力してやり遂げる" in joined
    assert "鉄棒に挑戦する姿が育った" in joined  # 保育の展開（development_notes）
    assert "就学先と見通しの持ち方を引き継ぐ" in joined  # 特に配慮すべき事項
    assert (
        "入園当初の不安から生き生きと表現する姿へ育った" in joined
    )  # 列4＝最終年度に至るまでの育ち


def test_fill_nursery_record_fills_name_and_school():
    tables = _tables(fill_docx("nursery_record", _nursery()))
    enroll = next(t for t in tables if "就学先" in "".join(c.text for r in t.rows for c in r.cells))
    joined = "".join(c.text for r in enroll.rows for c in r.cells)
    assert "かなたくん" in joined and "市立ひがし小学校" in joined


def test_fill_unknown_kind_raises():
    with pytest.raises(ValueError):
        fill_docx("diary", _record())


def test_fill_non_dict_raises():
    with pytest.raises(ValueError):
        fill_docx("child_record", "not a dict")  # type: ignore[arg-type]
