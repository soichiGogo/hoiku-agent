"""確定 entry（final_entry）→ 園の実 Word 様式（.docx）への流し込み（層A/web の presentation）。

設計コンテキスト §11（配信UI）／ §18（“園の様式で出す”最終形）。`chohyo_pdf.py`（ReportLab で
帳票PDF を描く）と対になる出力経路で、**園が実際に使っている Word フォーム（`web/templates/*.docx`）**を
そのまま雛形にして、確定 entry の値を該当セルへ流し込み、編集可能な .docx として返す。保育士は
Word で開いて微修正・印刷・PDF 化できる（＝Word が母艦の現場に寄せる）。**描画（流し込み）だけ**で、
必須欄・年齢分岐等の型の保証は harness が持つ（§5・ここは validation を持たない）。

依存は **python-docx（純 pip・システムライブラリ不要＝Dockerfile 不変）**。雛形 .docx は `web/templates/` に
同梱し実行時に外部取得しない（ローカル完結）。docx→PDF の**サーバ変換はしない**（LibreOffice 等の
重い依存を持ち込まない＝確定・綴じ用の PDF は `chohyo_pdf.py` が担い、Word 編集用は本モジュールが担う）。

配線済みスライス：**保育経過記録**＝5領域×子どもの姿（3–5）。**クラス月案（class_monthly）**＝
園フォーム（月間指導計画）の**全欄を直接埋める**（保育目標・先月の姿・区分×領域グリッド〔養護2本柱＋
教育5領域〕・食育/健康・安全/家庭/職員の連携・0–2 は個人目標小表を登場児ぶん）＝§18。評価系欄は月末記入＝
AI 非生成で空欄温存。**個別月案（monthly）**＝旧経路として温存＝個別月案の中身を「個人目標」小表へ写像し
クラス欄は保育士記入で温存（0-2 フォームのみ小表・3-5 はヘッダのみ）。**保育要録**＝公式様式
（こども家庭庁 保育所児童保育要録・参考例）の「保育に関する記録」の括弧ラベル（最終年度の重点/個人の重点/
保育の展開と子どもの育ち/特に配慮すべき事項）直下と列4（最終年度に至るまでの育ち）へ追記（ガイドラベルは残す）。
拡張は `_FILLERS` に kind を追加するだけ。
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.table import Table

from ..schemas import FiveDomains, ThreeViewpoint

_TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"


def _set_cell(cell, text: str) -> None:
    """セルの本文を text に置き換える（段落プロパティは残し、runs だけ差し替える＝書式温存寄り）。

    python-docx の `cell.text = ...` は段落・runs を丸ごと単一 run へ潰すため、既存の配置/フォント指定が
    失われやすい。ここでは最初の段落の runs をクリアして1つの run を足すに留め、空欄の書式を壊しにくくする。
    """
    para = cell.paragraphs[0]
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _find_table(doc: Document, *, header_contains: tuple[str, ...]) -> Table | None:
    """先頭行のセル文言に header_contains の語をすべて含む表を返す（署名でテンプレ表を同定）。

    セル位置のハードコードでなく見出し語で表を見つける＝テンプレの前後に表が増減しても壊れにくい。
    """
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        head = "".join(c.text for c in tbl.rows[0].cells)
        if all(tok in head for tok in header_contains):
            return tbl
    return None


def _find_table_by_any_cell(doc: Document, *tokens: str) -> Table | None:
    """全セルを走査し、tokens をすべて含む表を返す（見出しが row0 に無い様式用＝入所記録表 等）。"""
    for tbl in doc.tables:
        text = "".join(c.text for row in tbl.rows for c in row.cells)
        if all(tok in text for tok in tokens):
            return tbl
    return None


def _label_row_value_cell(tbl: Table, label: str):
    """表内で1列目が label のちょうどそのセル（＝流し込み先の2列目）を返す。無ければ None。

    保育経過記録の「領域×子どもの姿」表のように 1列目＝ラベル / 最終列＝記入 の様式を素直に扱う。
    """
    for row in tbl.rows:
        cells = row.cells
        if cells and cells[0].text.strip() == label:
            return cells[-1]
    return None


def _fiscal_year_label(period: str) -> str:
    """対象期間の自由記述先頭から年度（4月始まり）を推定して「YYYY年度」を返す。読めなければ空。"""
    m = re.search(r"(\d{4})\s*[-/年]\s*(\d{1,2})", str(period or ""))
    if not m:
        return ""
    year, month = int(m.group(1)), int(m.group(2))
    if not (1 <= month <= 12):
        return ""
    fiscal = year if month >= 4 else year - 1
    return f"{fiscal}年度"


_AGE_LABEL = {"0-2": "0〜2", "3-5": "3〜5"}


def _write_label_value_rows(tbl: Table, rows: list[tuple[str, str]]) -> None:
    """表のデータ行（見出し行を除く）を rows（ラベル, 内容）へ書き替える。行数は増減させる。

    テンプレのデータ行を先頭から再利用し、足りなければ `add_row` で増やし、余ったテンプレ行は削除する
    （枠組みが5領域より少ない 0-2＝3つの視点 等で空行が残らないようにする）。1列目＝ラベル / 最終列＝内容。
    """
    for i, (label, value) in enumerate(rows):
        if 1 + i < len(tbl.rows):
            cells = tbl.rows[1 + i].cells
        else:
            cells = tbl.add_row().cells
        _set_cell(cells[0], label)
        _set_cell(cells[-1], value)
    # 余ったテンプレ行を削除（先に list 化してから remove）。
    for extra in list(tbl.rows[1 + len(rows) :]):
        tbl._tbl.remove(extra._tr)


def _fill_domain_matrix(matrix: Table, entry: dict) -> None:
    """領域×子どもの姿の表を年齢別の枠組み＋「その他」で埋める（内容を落とさない）。

    テンプレの表は5領域行の固定様式。3-5 は枠組みが5領域で一致するので**テンプレの行順・書式を
    そのまま尊重**して該当ラベルへ流し込む（園様式の並びを崩さない）。0-2 は「3つの視点」タグで
    5領域行のどれにも載らない（旧実装では全部捨てられ空欄になっていた）ので、chohyo_pdf の年間
    マトリクスと同じく枠組み（3つの視点）へ**行を組み直す**。枠組みタグの無い叙述（10の姿のみ 等）は
    「その他」行へ受ける。
    """
    age_band = entry.get("age_band") or "3-5"
    framework = FiveDomains if age_band == "3-5" else ThreeViewpoint
    row_labels = [e.value for e in framework]
    by_label: dict[str, list[str]] = {label: [] for label in row_labels}
    other: list[str] = []
    for note in entry.get("development_notes") or []:
        note = note or {}
        desc = str(note.get("description") or "").strip()
        if not desc:
            continue
        tags = [str(t) for t in (note.get("tags") or [])]
        target = next((label for label in row_labels if label in tags), None)
        (by_label[target] if target else other).append(desc)

    template_labels = {matrix.rows[i].cells[0].text.strip() for i in range(1, len(matrix.rows))}
    if set(row_labels) <= template_labels:
        # 3-5：テンプレの行順・書式を保ち、ラベル一致セルへ流し込む（在庫の園様式を尊重）。
        for row in matrix.rows[1:]:
            label = row.cells[0].text.strip()
            if by_label.get(label):
                _set_cell(row.cells[-1], "\n".join(by_label[label]))
        if other:
            cells = matrix.add_row().cells
            _set_cell(cells[0], "その他")
            _set_cell(cells[-1], "\n".join(other))
    else:
        # 0-2：テンプレの5領域行は枠組みが違う。枠組み（3つの視点）＋その他へ行を組み直す。
        out_rows = [(label, "\n".join(by_label[label])) for label in row_labels]
        if other:
            out_rows.append(("その他", "\n".join(other)))
        _write_label_value_rows(matrix, out_rows)


def _fill_child_record(entry: dict) -> Document:
    """保育経過記録フォームへ確定 entry を流し込む。

    - ヘッダ表：児童名・年度（対象期間から推定）・歳児（年齢帯）。生年月日・担任印は手書き欄のまま。
    - 対象期間表：period をそのまま。
    - 領域×子どもの姿表：development_notes を年齢別の枠組み（0-2＝3つの視点／3-5＝5領域）で各行に
      振り分ける（複数は改行連結）。枠組みタグを持たない叙述（10の姿のみ 等）は「その他」行へ受けて
      内容を落とさない（chohyo_pdf の年間マトリクスと同じ扱い）。
    """
    doc = Document(str(_TEMPLATE_DIR / "child_record.docx"))

    # ── ヘッダ（年度／歳児／児童名） ──
    header = _find_table(doc, header_contains=("年度", "歳児", "担任"))
    if header is not None:
        cells = header.rows[0].cells
        # r0: [年度, <値>, 歳児, <値>, 担任, <値>, 印]
        _set_cell(cells[1], _fiscal_year_label(str(entry.get("period") or "")))
        _set_cell(cells[3], _AGE_LABEL.get(entry.get("age_band") or "3-5", ""))
        # r1: [児童名, <値(結合)>, …, 生年月日, …]。最終列が生年月日の手書き欄なので _label_row_value_cell は
        # 使わず（それは最終列を返す）、児童名ラベル直後の結合セル（index1）へ入れる。
        if len(header.rows) > 1:
            _set_cell(header.rows[1].cells[1], str(entry.get("child_id") or ""))

    # ── 対象期間 ──
    period_tbl = _find_table(doc, header_contains=("対象期間",))
    if period_tbl is not None:
        cell = _label_row_value_cell(period_tbl, "対象期間")
        if cell is not None:
            _set_cell(cell, str(entry.get("period") or ""))

    # ── 領域×子どもの姿（AI 内容の本体・年齢別枠組み＋その他で内容を落とさない） ──
    matrix = _find_table(doc, header_contains=("領域", "子どもの姿"))
    if matrix is not None:
        _fill_domain_matrix(matrix, entry)

    return doc


def _monthly_label(month: str) -> str:
    """月案の対象月（YYYY-MM）を「YYYY年度 M月」へ整形する（4月始まりの年度）。読めなければ原文。"""
    m = re.match(r"\s*(\d{4})\s*[-/年]\s*(\d{1,2})", str(month or ""))
    if not m:
        return str(month or "")
    year, mm = int(m.group(1)), int(m.group(2))
    if not (1 <= mm <= 12):
        return str(month or "")
    fiscal = year if mm >= 4 else year - 1
    return f"{fiscal}年度 {mm}月"


def _monthly_individual_summary(entry: dict) -> str:
    """個別月案の中身を「個人目標」小表の『ねらい・配慮』1セルへ収める要約（AI 内容を落とさない）。

    園フォームの個人目標欄は 氏名/子どもの姿/ねらい・配慮/評価・反省 の4列しかないため、我々の
    養護2本柱・教育ねらい・環境援助を『ねらい・配慮』に構造化して束ねる（クラス欄は保育士記入で温存）。
    """
    parts: list[str] = []
    goals = str(entry.get("monthly_goals") or "").strip()
    if goals:
        parts.append(f"今月のねらい：{goals}")
    nlife = str(entry.get("nurturing_life") or "").strip()
    nemo = str(entry.get("nurturing_emotion") or "").strip()
    if nlife:
        parts.append(f"養護（生命の保持）：{nlife}")
    if nemo:
        parts.append(f"養護（情緒の安定）：{nemo}")
    edu_aims = [
        str((note or {}).get("aim") or "").strip()
        for note in entry.get("education") or []
        if str((note or {}).get("aim") or "").strip()
    ]
    if edu_aims:
        parts.append("教育：" + "／".join(edu_aims))
    env = str(entry.get("environment_support") or "").strip()
    if env:
        parts.append(f"環境・援助：{env}")
    return "\n".join(parts)


def _fill_monthly(entry: dict) -> Document:
    """個別月案を園の月案フォーム（クラス月案）へ流し込む。

    園フォームは**クラス月案**（区分×領域グリッド＝クラス全体のねらい）で、我々の出力は**個別月案**。
    決定（2026-07-05）に従い、**個別の中身は「個人目標」小表へ写像**し、クラス全体欄（保育目標・
    区分×領域グリッド・食育・評価等）は当面**保育士記入で空欄温存**する（勝手に埋めない）。

    - ヘッダ：年度・月／年齢帯（機械メタなので埋める）。
    - **個人目標小表（0-2 フォームのみ存在）**：氏名（月齢）／子どもの姿（前月の姿）／ねらい・配慮
      （今月のねらい＋養護2本柱＋教育＋環境援助の要約）／評価・反省 の1行に写像。
    - 3-5 フォームは個人目標小表が無い（純クラス様式）ため、ヘッダのみ流し込む＝個別内容の置き場が
      様式に無い（クラス月案スキーマ対応は後続・§18）。
    """
    age = entry.get("age_band") or "0-2"
    template = "monthly_0_2.docx" if age == "0-2" else "monthly_3_5.docx"
    doc = Document(str(_TEMPLATE_DIR / template))

    # ── ヘッダ（年度・月／年齢） ──
    header = _find_table(doc, header_contains=("年度・月", "クラス"))
    if header is not None:
        _set_cell(header.rows[0].cells[1], _monthly_label(str(entry.get("month") or "")))
        if len(header.rows) > 1:
            _set_cell(header.rows[1].cells[1], f"{_AGE_LABEL.get(age, age)}歳児（　　名）")

    # ── 個人目標小表（0-2 のみ）：個別月案の中身を1行へ写像 ──
    goals_tbl = _find_table(doc, header_contains=("個人目標",))
    if goals_tbl is not None:
        # r0=見出し「個人目標…」、r1=列見出し（氏名/子どもの姿/ねらい・配慮/評価・反省）、r2 以降=記入行。
        row = goals_tbl.rows[2] if len(goals_tbl.rows) > 2 else None
        if row is not None and len(row.cells) >= 4:
            name = str(entry.get("child_id") or "")
            months = str(entry.get("age_months") or "").strip()
            _set_cell(row.cells[0], f"{name}（{months}）" if months else name)
            _set_cell(row.cells[1], str(entry.get("prev_child_state") or ""))
            _set_cell(row.cells[2], _monthly_individual_summary(entry))
            _set_cell(row.cells[3], str(entry.get("evaluation_reflection") or ""))

    return doc


def _set_after_label(tbl: Table, label: str, value: str) -> bool:
    """表内で text が label のセルの**次のセル**（同じ行の隣）に value を入れる（label|value 横並び用）。

    _label_row_value_cell（行の最終列を返す）と違い、ラベルの直後セルに入れる＝食育|健康・安全 のように
    1行に label|value|label|value が並ぶ様式に使う。見つけて入れたら True。
    """
    for row in tbl.rows:
        cells = row.cells
        for i, cell in enumerate(cells[:-1]):
            if cell.text.strip() == label:
                _set_cell(cells[i + 1], value)
                return True
    return False


def _fill_class_monthly(entry: dict) -> Document:
    """クラス月案（園の実様式＝月間指導計画）フォームへ確定 entry を流し込む（§18）。

    園フォーム（monthly_0_2.docx / monthly_3_5.docx＝A4 横）の欄構成そのものに埋める：
    - ヘッダ表：年度・月／クラス／年齢（担任・園長・主任は押印欄＝手書きのまま）。
    - 上部の単欄表：今月の保育目標／先月の子どもの姿／今月の行事／保護者支援。
    - 区分×領域グリッド：領域名で行を引き当て、ねらい／環境・構成／子どもの姿／援助・配慮を埋める。
    - 連携表：食育／健康・安全／家庭との連携／職員間の連携。
    - 個人目標小表（0–2 フォームのみ存在）：登場児ぶんを行に埋める（不足行は add_row）。
    評価系欄（保育者の評価等）は月末記入＝AI 非生成なので触らない（空欄温存）。描画のみ（型の保証は harness）。
    """
    age = entry.get("age_band") or "0-2"
    template = "monthly_0_2.docx" if age == "0-2" else "monthly_3_5.docx"
    doc = Document(str(_TEMPLATE_DIR / template))

    # ── ヘッダ（年度・月／クラス／年齢） ──
    header = _find_table(doc, header_contains=("年度・月", "クラス", "担任"))
    if header is not None:
        _set_cell(header.rows[0].cells[1], str(entry.get("month") or ""))
        _set_cell(header.rows[0].cells[3], str(entry.get("class_name") or ""))
        if len(header.rows) > 1:
            _set_cell(header.rows[1].cells[1], f"{_AGE_LABEL.get(age, age)}歳児")

    # ── 上部の単欄（今月の保育目標／先月の子どもの姿／今月の行事／保護者支援） ──
    top = _find_table(doc, header_contains=("今月の保育目標",))
    if top is not None:
        for label, key in (
            ("今月の保育目標", "monthly_goal"),
            ("先月の子どもの姿", "prev_month_state"),
            ("今月の行事", "events"),
            ("保護者支援", "parent_support"),
        ):
            cell = _label_row_value_cell(top, label)
            if cell is not None:
                _set_cell(cell, str(entry.get(key) or ""))

    # ── 区分×領域グリッド（領域名で行を引き当て、内容4欄を埋める） ──
    grid = _find_table(doc, header_contains=("区分", "領域", "ねらい"))
    if grid is not None:
        by_domain = {
            str((r or {}).get("domain") or "").strip(): (r or {}) for r in entry.get("grid") or []
        }
        for row in grid.rows[1:]:  # 見出し行を除く
            cells = row.cells
            if len(cells) < 6:
                continue
            domain = cells[1].text.strip()
            src = by_domain.get(domain)
            if not src:
                continue
            _set_cell(cells[2], str(src.get("aim") or ""))
            _set_cell(cells[3], str(src.get("environment") or ""))
            _set_cell(cells[4], str(src.get("child_state") or ""))
            _set_cell(cells[5], str(src.get("support") or ""))

    # ── 連携（食育／健康・安全／家庭との連携／職員間の連携＝label|value 横並び） ──
    liaison = _find_table(doc, header_contains=("食育", "健康・安全"))
    if liaison is not None:
        _set_after_label(liaison, "食育", str(entry.get("syokuiku") or ""))
        _set_after_label(liaison, "健康・安全", str(entry.get("health_safety") or ""))
        _set_after_label(liaison, "家庭との連携", str(entry.get("family_liaison") or ""))
        _set_after_label(liaison, "職員間の連携", str(entry.get("staff_liaison") or ""))

    # ── 個人目標小表（0–2 フォームのみ）：登場児ぶんを行に埋める（不足行は add_row） ──
    goals = entry.get("individual_goals") or []
    goals_tbl = _find_table(doc, header_contains=("個人目標",))
    if goals_tbl is not None and goals:
        # r0=見出し「個人目標…」、r1=列見出し（氏名/子どもの姿/ねらい・配慮/評価・反省）、r2 以降=記入行。
        data_start = 2
        for i, goal in enumerate(goals):
            goal = goal or {}
            ri = data_start + i
            row = goals_tbl.rows[ri] if ri < len(goals_tbl.rows) else goals_tbl.add_row()
            if len(row.cells) < 4:
                continue
            name = str(goal.get("child_id") or "")
            months = str(goal.get("age_months") or "").strip()
            _set_cell(row.cells[0], f"{name}（{months}）" if months else name)
            _set_cell(row.cells[1], str(goal.get("child_state") or ""))
            _set_cell(row.cells[2], str(goal.get("aim_support") or ""))
            # 評価・反省（col3）は月末記入＝空欄温存。

    return doc


def _append_cell_lines(cell, lines: list[str]) -> None:
    """セルの既存内容（括弧ラベル等）を残したまま、内容を段落として下に追記する。

    `_set_cell`（置換）と違い、公式様式の（最終年度の重点）等のガイドラベルを消さずに内容を足す。
    空行は追記しない。
    """
    for line in lines:
        text = str(line).strip()
        if not text:
            continue
        cell.add_paragraph().add_run(text)


def _fill_nursery_record(entry: dict) -> Document:
    """保育要録（保育所児童保育要録・保育に関する記録）を公式様式（こども家庭庁 参考例）へ流し込む。

    様式は「入所に関する記録」（原簿系＝AI外・手書き）と「保育に関する記録」（AI 生成部）の2部。
    後者の列3「保育の過程と子どもの育ちに関する事項」に（最終年度の重点）（個人の重点）
    （保育の展開と子どもの育ち）（特に配慮すべき事項）が段組みで並び、列4に「最終年度に至るまでの
    育ち」が入る。**各括弧ラベルの直下に内容を追記**し（ラベルは残す＝様式のガイドを壊さない）、
    列4に growth を追記する。氏名・就学先だけ入所記録表に添える（生年月日・押印等は手書き欄のまま）。
    5領域のねらい参考列・別紙（10の姿）は様式の固定参照＝触らない。描画のみ（型の保証は harness＝§5）。
    """
    doc = Document(str(_TEMPLATE_DIR / "nursery_record.docx"))

    # ── 入所に関する記録（原簿系）：氏名・就学先だけ添える（他は手書き。見出しが row0 に無い様式なので
    #    全セル走査で表を同定する） ──
    enroll = _find_table_by_any_cell(doc, "就学先", "保護者")
    if enroll is not None:
        for row in enroll.rows:
            cells = row.cells
            head = cells[0].text.replace("　", "").replace(" ", "").strip()
            label1 = (
                cells[1].text.replace("　", "").replace(" ", "").strip() if len(cells) > 1 else ""
            )
            if head.startswith("児") and "氏名" in label1 and len(cells) > 3:
                _set_cell(cells[3], str(entry.get("child_id") or ""))
            elif head.startswith("就学先") and len(cells) > 1:
                school = str(entry.get("school_name") or "").strip()
                if school:
                    _set_cell(cells[1], school)

    # ── 保育に関する記録（AI 生成部）：括弧ラベル直下へ内容を追記 ──
    rec = _find_table(doc, header_contains=("保育の過程と子どもの育ちに関する事項",))
    if rec is None:
        return doc
    dev_lines = []
    for note in entry.get("development_notes") or []:
        note = note or {}
        desc = str(note.get("description") or "").strip()
        if not desc:
            continue
        tags = "・".join(str(t) for t in (note.get("tags") or []))
        dev_lines.append(f"{desc}（{tags}）" if tags else desc)
    # 括弧ラベル → 追記する内容（列3）。各ラベルの最初の出現セルにだけ追記する。
    col3_content = {
        "（最終年度の重点）": [str(entry.get("final_year_focus") or "")],
        "（個人の重点）": [str(entry.get("individual_focus") or "")],
        "（保育の展開と子どもの育ち）": dev_lines,
        "（特に配慮すべき事項）": [str(entry.get("special_notes") or "")],
    }
    done: set[str] = set()
    growth_done = False
    for row in rec.rows:
        cells = row.cells
        if len(cells) > 3:
            label = cells[3].text.strip().split("\n")[0].strip()
            if label in col3_content and label not in done:
                _append_cell_lines(cells[3], col3_content[label])
                done.add(label)
        # 列4（最終年度に至るまでの育ち）：見出し行を除く最初の内容セルへ growth を追記。
        if len(cells) > 4 and not growth_done:
            head4 = cells[4].text.strip()
            if head4 == "" or head4 == "最終年度に至るまでの育ちに関する事項":
                if head4 == "":
                    _append_cell_lines(cells[4], [str(entry.get("growth_until_final") or "")])
                    growth_done = True

    return doc


# kind → filler。新しい書類は filler を足すだけで拡張できる。
_FILLERS = {
    "child_record": _fill_child_record,
    "monthly": _fill_monthly,
    "class_monthly": _fill_class_monthly,
    "nursery_record": _fill_nursery_record,
}


def supported_kinds() -> list[str]:
    """docx 流し込みに対応済みの kind 一覧（UI が Word ダウンロードボタンの出し分けに使う）。"""
    return list(_FILLERS)


def fill_docx(kind: str, entry: dict) -> bytes:
    """確定 entry（dict）を園の実 Word 様式へ流し込み、.docx の bytes を返す。

    未対応 kind・entry が dict でない場合は ValueError（route が 400 で可視化・握りつぶさない）。
    描画のみ（型検査はしない＝空欄は空セルのまま。型の保証は harness＝§5）。
    """
    filler = _FILLERS.get(kind)
    if filler is None:
        raise ValueError(f"docx 未対応の kind: {kind!r}（対応: {', '.join(_FILLERS) or 'なし'}）")
    if not isinstance(entry, dict):
        raise ValueError("entry は dict である必要があります")
    doc = filler(entry)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
