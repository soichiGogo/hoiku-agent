"""アップロードされたファイル（bytes）→ LLM 入力コンテンツへの**決定的**変換（§11 presentation・IO）。

「書類を見る」タブのアップロード取込（`/api/parse-upload`）の前段。責務は**フォーマット変換だけ**で、
中身の解釈（スキーマへの写像）は agents/ の LLM（`upload_parser_agent`）が担う（決定ロジックは持たない）。
`chohyo_pdf.py`（entry→PDF）・`docx_fill.py`（entry→docx）と同じ「web に置く純粋なフォーマット変換」の仲間。

方針:
- **docx** … python-docx で段落＋表セルのテキストを抽出（既存依存）。
- **xlsx** … openpyxl でシートのセル値を行ごとに抽出（`data_only` で数式は計算値）。
- **pdf** … テキスト抽出せず、**Gemini マルチモーダルへ生 bytes を渡す**（`inline_data`＝
  application/pdf）。レイアウト/表/手書きスキャンをそのままモデルに読ませる（追加のPDFライブラリを持たない）。

抽出テキストはトークン/メモリを無制限に食わないよう上限で切り詰める（`MAX_TEXT_CHARS`）。
旧バイナリ様式（.doc/.xls）は openxml でないため未対応＝正直に ValueError（握りつぶさない）。
"""

from __future__ import annotations

import io
from dataclasses import dataclass

# 拡張子 → 論理フォーマット（ブラウザの MIME は空/不正がありうるので拡張子を第一に見る）。
_EXT_FORMAT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
}
# MIME → 論理フォーマット（拡張子で判別できないときのフォールバック）。
_MIME_FORMAT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
}

SUPPORTED_FORMATS = ("pdf", "docx", "xlsx")

# 抽出テキストの上限（docx/xlsx）。超過分は切り詰める（トークン/メモリのガード）。
MAX_TEXT_CHARS = 60_000
# アップロード全体の上限バイト数（ルート側でも弾くが、抽出側の安全網）。
MAX_UPLOAD_BYTES = 25 * 1024 * 1024


@dataclass
class ExtractedUpload:
    """アップロード1件を LLM に渡せる形へ正規化した結果（決定的）。

    - `fmt`＝"pdf"/"docx"/"xlsx"。
    - `text`＝docx/xlsx から抽出した本文テキスト（pdf は空）。
    - `pdf_bytes`＝pdf のときだけ Gemini へ渡す生 bytes（それ以外は None）。
    """

    fmt: str
    text: str = ""
    pdf_bytes: bytes | None = None


def detect_format(filename: str, mime_type: str | None) -> str | None:
    """ファイル名の拡張子（優先）→ MIME で論理フォーマットを決める。判別不能は None。"""
    name = (filename or "").lower().strip()
    for ext, fmt in _EXT_FORMAT.items():
        if name.endswith(ext):
            return fmt
    return _MIME_FORMAT.get((mime_type or "").strip().lower())


def extract_upload(filename: str, mime_type: str | None, data: bytes) -> ExtractedUpload:
    """アップロード bytes を `ExtractedUpload` へ変換する（決定的）。

    未対応フォーマット・空データ・過大サイズ・壊れたファイルは ValueError（正直に上げる＝
    ルートが 400 に変換。黙って空 entry を作らない）。
    """
    if not data:
        raise ValueError("ファイルが空です")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(f"ファイルが大きすぎます（上限 {MAX_UPLOAD_BYTES // (1024 * 1024)}MB）")
    fmt = detect_format(filename, mime_type)
    if fmt is None:
        raise ValueError(
            "対応していない形式です（pdf / docx / xlsx のみ。.doc / .xls の旧形式は不可）"
        )
    if fmt == "pdf":
        return ExtractedUpload(fmt="pdf", pdf_bytes=data)
    if fmt == "docx":
        return ExtractedUpload(fmt="docx", text=_truncate(_docx_to_text(data)))
    if fmt == "xlsx":
        return ExtractedUpload(fmt="xlsx", text=_truncate(_xlsx_to_text(data)))
    raise ValueError(f"未対応の形式: {fmt}")  # 到達しない（SUPPORTED_FORMATS と同期）


def to_parts(extracted: ExtractedUpload):
    """`ExtractedUpload` を Gemini（ADK LlmAgent）の入力 Part 列へ変換する。

    genai types への依存はここに閉じ込める（`extract_upload`＝テキスト抽出は genai 非依存で単体テスト可能）。
    pdf は `inline_data`（マルチモーダル）で生 bytes、docx/xlsx は抽出テキストを text Part にする。
    """
    from google.genai import types

    if extracted.fmt == "pdf" and extracted.pdf_bytes is not None:
        return [
            types.Part(
                inline_data=types.Blob(mime_type="application/pdf", data=extracted.pdf_bytes)
            )
        ]
    return [types.Part(text=extracted.text)]


def _truncate(text: str) -> str:
    if len(text) <= MAX_TEXT_CHARS:
        return text
    return text[:MAX_TEXT_CHARS] + "\n…（以下省略：文字数上限に達したため切り詰め）"


def _join_bounded(lines) -> str:
    """行イテレータを結合するが、累積が MAX_TEXT_CHARS を超えたら打ち切る（zip 爆弾の展開コストを有界化）。

    _truncate は「全部蓄積してから末尾を切る」ため、高圧縮 docx/xlsx（展開後は膨大なセル/段落）だと
    切り詰め前に list が非有界に膨れて OOM しうる。ここで蓄積側を止め、消費を打ち切ることで
    openpyxl の read_only ストリームや python-docx のイテレーションもそれ以上読まない。
    """
    out: list[str] = []
    total = 0
    for line in lines:
        out.append(line)
        total += len(line) + 1
        if total > MAX_TEXT_CHARS:
            break
    return _truncate("\n".join(out))


def _iter_docx_lines(doc):
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            yield t
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                yield " | ".join(cells)


def _docx_to_text(data: bytes) -> str:
    """docx の段落＋表セルをテキスト化する（python-docx）。壊れは ValueError。"""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001  壊れた docx/非 openxml は正直に上げる
        raise ValueError(f"docx を読み取れませんでした: {e}") from e
    return _join_bounded(_iter_docx_lines(doc))


def _iter_xlsx_lines(wb):
    for ws in wb.worksheets:
        yield f"# シート: {ws.title}"
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if any(cells):
                yield " | ".join(cells).rstrip(" |")


def _xlsx_to_text(data: bytes) -> str:
    """xlsx の各シートのセル値を行ごとにテキスト化する（openpyxl・数式は計算値）。壊れは ValueError。"""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001  壊れた xlsx/非 openxml は正直に上げる
        raise ValueError(f"xlsx を読み取れませんでした: {e}") from e
    try:
        return _join_bounded(_iter_xlsx_lines(wb))
    finally:
        wb.close()
